import json
import re
import statistics
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

ROOT_DIR = Path(__file__).resolve().parents[2]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from api.main import create_app
from scripts.day2.create_sample_pdf import build_pdf_bytes
from src.embeddings import DeterministicEmbeddingModel
from src.llm import ChatModel


CORPUS_DIR = Path("data/day12_corpus")
REPORT_DIR = Path("docs/day12")
REPORT_PATH = REPORT_DIR / "system_integration_report.md"
RESULTS_PATH = REPORT_DIR / "system_integration_results.json"


@dataclass(frozen=True)
class CorpusFile:
    file_name: str
    file_type: str
    content: str


@dataclass(frozen=True)
class EvalCase:
    question: str
    expected_file: str
    expected_keywords: list[str]
    scenario: str


class ExtractiveTestChatModel(ChatModel):
    """Deterministic local model for repeatable integration tests."""

    def generate(self, prompt: str) -> str:
        question = _extract_question(prompt)
        context = _extract_between(prompt, "銆愭绱㈣祫鏂欍€?", "銆愮敤鎴烽棶棰樸€?") or prompt
        if question in context:
            context = context.split(question, 1)[0]
        query_tokens = _content_tokens(question)
        query_phrases = _phrases(query_tokens)
        best_sentence = ""
        best_score = -1.0
        for sentence in _split_sentences(context):
            if _is_prompt_marker(sentence):
                continue
            sentence_tokens = _content_tokens(sentence)
            sentence_phrases = _phrases(sentence_tokens)
            token_overlap = len(set(query_tokens) & set(sentence_tokens))
            phrase_overlap = len(query_phrases & sentence_phrases)
            number_overlap = len(set(_numbers(question)) & set(_numbers(sentence)))
            score = token_overlap + phrase_overlap * 2.5 + number_overlap * 3.0
            if score > best_score:
                best_sentence = sentence
                best_score = score
        if not best_sentence or best_score <= 0:
            return "No relevant information was found in the uploaded documents."
        return best_sentence.strip()


def build_corpus() -> list[CorpusFile]:
    return [
        CorpusFile("leave_policy.txt", "policy", "Annual leave policy. Full-time employees receive 15 annual leave days each calendar year. Leave requests must be submitted at least 3 business days before the planned absence. Unused annual leave cannot exceed 5 carryover days."),
        CorpusFile("expense_policy.docx", "policy", "Expense reimbursement policy. Taxi receipts must be submitted within 30 days. Hotel reimbursement requires an invoice and manager approval. Meal reimbursement is capped at 120 CNY per person per day."),
        CorpusFile("product_manual.pdf", "product manual", "Product Atlas manual. The Atlas device supports offline sync every 10 minutes. The default admin password must be changed during first login. Firmware updates are installed from the system settings page."),
        CorpusFile("vpn_faq.txt", "FAQ", "VPN FAQ. If VPN login fails, reset the MFA token and retry with the company account. VPN access is disabled automatically after 45 days of inactivity. Remote employees should use the Asia gateway for best latency."),
        CorpusFile("security_policy.docx", "policy", "Security policy. Passwords must be at least 12 characters and include two character classes. Sensitive files must be stored only in the approved document vault. USB storage devices are prohibited unless approved by IT security."),
        CorpusFile("support_faq.pdf", "FAQ", "Customer support FAQ. Priority P1 tickets require the first response within 15 minutes. Refund requests are handled by the billing team. Support agents must attach the troubleshooting log before escalation."),
        CorpusFile("onboarding_guide.txt", "guide", "Onboarding guide. New hires must complete compliance training within 7 days. The HR portal contains payroll setup and emergency contact forms. A mentor is assigned during the first week."),
        CorpusFile("sla_policy.docx", "policy", "Service level policy. Standard incidents have a target resolution time of 2 business days. Enterprise customers receive 99.9 percent monthly availability. Maintenance windows are announced at least 72 hours in advance."),
    ]


def build_eval_cases() -> list[EvalCase]:
    return [
        EvalCase("How many annual leave days do full-time employees receive?", "leave_policy.txt", ["15", "annual leave"], "HR policy"),
        EvalCase("How early must leave requests be submitted?", "leave_policy.txt", ["3 business days"], "HR policy"),
        EvalCase("How many annual leave days can be carried over?", "leave_policy.txt", ["5 carryover days"], "HR policy"),
        EvalCase("When must taxi receipts be submitted?", "expense_policy.docx", ["30 days"], "Finance policy"),
        EvalCase("What is required for hotel reimbursement?", "expense_policy.docx", ["invoice", "manager approval"], "Finance policy"),
        EvalCase("What is the meal reimbursement cap?", "expense_policy.docx", ["120 CNY"], "Finance policy"),
        EvalCase("How often does the Atlas device perform offline sync?", "product_manual.pdf", ["10 minutes"], "Product manual"),
        EvalCase("What must happen to the default admin password?", "product_manual.pdf", ["changed", "first login"], "Product manual"),
        EvalCase("Where are firmware updates installed from?", "product_manual.pdf", ["system settings"], "Product manual"),
        EvalCase("What should an employee do if VPN login fails?", "vpn_faq.txt", ["reset", "MFA token"], "IT FAQ"),
        EvalCase("When is VPN access automatically disabled?", "vpn_faq.txt", ["45 days"], "IT FAQ"),
        EvalCase("Which gateway should remote employees use?", "vpn_faq.txt", ["Asia gateway"], "IT FAQ"),
        EvalCase("What is the minimum password length?", "security_policy.docx", ["12 characters"], "Security policy"),
        EvalCase("Where must sensitive files be stored?", "security_policy.docx", ["document vault"], "Security policy"),
        EvalCase("Are USB storage devices allowed?", "security_policy.docx", ["prohibited", "IT security"], "Security policy"),
        EvalCase("What is the first response target for P1 tickets?", "support_faq.pdf", ["15 minutes"], "Support FAQ"),
        EvalCase("Who handles refund requests?", "support_faq.pdf", ["billing team"], "Support FAQ"),
        EvalCase("When must new hires complete compliance training?", "onboarding_guide.txt", ["7 days"], "Onboarding"),
        EvalCase("What is the target resolution time for standard incidents?", "sla_policy.docx", ["2 business days"], "SLA"),
        EvalCase("How far ahead are maintenance windows announced?", "sla_policy.docx", ["72 hours"], "SLA"),
    ]


def write_corpus_files(corpus: list[CorpusFile]) -> list[Path]:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    for item in corpus:
        path = CORPUS_DIR / item.file_name
        if path.suffix == ".txt":
            path.write_text(item.content, encoding="utf-8")
        elif path.suffix == ".docx":
            doc = DocxDocument()
            doc.add_heading(item.file_name, level=1)
            for sentence in _split_sentences(item.content):
                doc.add_paragraph(sentence)
            doc.save(path)
        elif path.suffix == ".pdf":
            path.write_bytes(build_pdf_bytes([item.file_name, *_split_sentences(item.content)]))
        else:
            raise ValueError(f"Unsupported corpus file: {path}")
        paths.append(path)
    return paths


def run_integration_test() -> dict[str, Any]:
    started_at = time.perf_counter()
    file_paths = write_corpus_files(build_corpus())
    eval_cases = build_eval_cases()
    app = create_app(
        embedding_model=DeterministicEmbeddingModel(dimensions=128),
        chat_model=ExtractiveTestChatModel(),
        collection_name=f"day12_integration_{int(time.time() * 1000)}",
    )
    client = TestClient(app)

    upload_results = []
    for path in file_paths:
        upload_start = time.perf_counter()
        with path.open("rb") as file_obj:
            response = client.post(
                "/documents/upload",
                files={"file": (path.name, file_obj, _content_type_for(path))},
            )
        elapsed_ms = round((time.perf_counter() - upload_start) * 1000, 2)
        response.raise_for_status()
        payload = response.json()
        payload["upload_ms"] = elapsed_ms
        upload_results.append(payload)

    qa_results = []
    for case in eval_cases:
        ask_start = time.perf_counter()
        response = client.post("/qa/ask", json={"question": case.question})
        elapsed_ms = round((time.perf_counter() - ask_start) * 1000, 2)
        response.raise_for_status()
        payload = response.json()
        source_files = [source["file_name"] for source in payload["sources"]]
        answer = payload["answer"]
        qa_results.append(
            {
                "scenario": case.scenario,
                "question": case.question,
                "expected_file": case.expected_file,
                "expected_keywords": case.expected_keywords,
                "answer": answer,
                "source_files": source_files,
                "response_ms": elapsed_ms,
                "retrieval_hit": case.expected_file in source_files,
                "answer_hit": all(keyword.lower() in answer.lower() for keyword in case.expected_keywords),
            }
        )

    results = {
        "summary": _calculate_metrics(upload_results, qa_results, started_at),
        "uploaded_documents": upload_results,
        "qa_results": qa_results,
    }
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    RESULTS_PATH.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
    REPORT_PATH.write_text(_render_report(results), encoding="utf-8")
    return results


def _calculate_metrics(upload_results: list[dict[str, Any]], qa_results: list[dict[str, Any]], started_at: float) -> dict[str, Any]:
    question_count = len(qa_results)
    response_times = [item["response_ms"] for item in qa_results]
    upload_times = [item["upload_ms"] for item in upload_results]
    retrieval_accuracy = sum(1 for item in qa_results if item["retrieval_hit"]) / question_count
    answer_accuracy = sum(1 for item in qa_results if item["answer_hit"]) / question_count
    p95_response = _percentile(response_times, 95)
    return {
        "document_count": len(upload_results),
        "question_count": question_count,
        "retrieval_accuracy": round(retrieval_accuracy, 4),
        "answer_accuracy": round(answer_accuracy, 4),
        "avg_response_ms": round(statistics.mean(response_times), 2),
        "p95_response_ms": round(p95_response, 2),
        "max_response_ms": round(max(response_times), 2),
        "avg_upload_ms": round(statistics.mean(upload_times), 2),
        "total_runtime_ms": round((time.perf_counter() - started_at) * 1000, 2),
        "performance_requirement": "PASS" if p95_response < 1000 else "REVIEW",
        "notes": [
            "Uses local deterministic embeddings and a deterministic extractive chat model.",
            "Validates upload, parsing, chunking, vector write/search, answer generation, citation, and reporting.",
        ],
    }


def _render_report(results: dict[str, Any]) -> str:
    summary = results["summary"]
    lines = [
        "# Day 12 System Integration Test Report",
        "",
        "## Scope",
        "",
        "This report validates the complete offline API flow: upload documents, parse, split, store vectors, ask questions, return answers, and cite sources.",
        "",
        "## Summary",
        "",
        f"- Uploaded documents: {summary['document_count']}",
        f"- Evaluation questions: {summary['question_count']}",
        f"- Retrieval accuracy: {summary['retrieval_accuracy']:.2%}",
        f"- Answer accuracy: {summary['answer_accuracy']:.2%}",
        f"- Average QA response time: {summary['avg_response_ms']} ms",
        f"- P95 QA response time: {summary['p95_response_ms']} ms",
        f"- Max QA response time: {summary['max_response_ms']} ms",
        f"- Average upload time: {summary['avg_upload_ms']} ms",
        f"- Performance requirement: {summary['performance_requirement']}",
        "",
        "## Uploaded Documents",
        "",
        "| File | Chunks | Vectors | Upload ms |",
        "| --- | ---: | ---: | ---: |",
    ]
    for item in results["uploaded_documents"]:
        lines.append(f"| {item['file_name']} | {item['chunk_count']} | {item['vector_count']} | {item['upload_ms']} |")
    lines.extend(["", "## Question Results", "", "| # | Scenario | Expected source | Retrieval | Answer | Response ms |", "| ---: | --- | --- | --- | --- | ---: |"])
    for index, item in enumerate(results["qa_results"], start=1):
        retrieval = "PASS" if item["retrieval_hit"] else "FAIL"
        answer = "PASS" if item["answer_hit"] else "FAIL"
        lines.append(f"| {index} | {item['scenario']} | {item['expected_file']} | {retrieval} | {answer} | {item['response_ms']} |")
    lines.extend(
        [
            "",
            "## Issues And Solutions",
            "",
            "- External LLM calls were replaced with a deterministic local model for repeatable integration testing.",
            "- The test corpus is ASCII English to avoid Windows console encoding problems during automated verification.",
            "- Chroma telemetry warnings may appear in the console, but they do not affect test results.",
            "",
            "## Conclusion",
            "",
            "The core technical choices are feasible when retrieval and answer accuracy meet the thresholds in the automated test.",
        ]
    )
    return "\n".join(lines) + "\n"


def _split_sentences(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+", text.strip()) if part.strip()]


def _tokens(text: str) -> list[str]:
    return re.findall(r"[a-zA-Z0-9]+", text.lower())


STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "can",
    "do",
    "does",
    "for",
    "from",
    "how",
    "if",
    "in",
    "is",
    "it",
    "many",
    "must",
    "of",
    "on",
    "or",
    "should",
    "the",
    "to",
    "what",
    "when",
    "where",
    "which",
    "who",
    "with",
}


def _content_tokens(text: str) -> list[str]:
    return [token for token in _tokens(text) if token not in STOPWORDS and len(token) > 1]


def _phrases(tokens: list[str]) -> set[str]:
    return {f"{tokens[index]} {tokens[index + 1]}" for index in range(len(tokens) - 1)}


def _numbers(text: str) -> list[str]:
    return re.findall(r"\d+(?:\.\d+)?", text)


def _is_prompt_marker(text: str) -> bool:
    marker_fragments = ["銆", "棶棰", "绛", "資料", "璧勬枡"]
    return any(fragment in text for fragment in marker_fragments) and not re.search(r"[a-zA-Z]{3,}", text)


def _extract_between(text: str, start_marker: str, end_marker: str) -> str:
    if start_marker not in text:
        return ""
    rest = text.split(start_marker, 1)[1]
    if end_marker in rest:
        rest = rest.split(end_marker, 1)[0]
    return rest.strip()


def _extract_question(text: str) -> str:
    marked_question = _extract_between(text, "銆愮敤鎴烽棶棰樸€?", "銆愬洖绛斻€?")
    if marked_question and "?" in marked_question:
        return marked_question
    questions = re.findall(r"([A-Z][^\n?]{8,}\?)", text)
    return questions[-1].strip() if questions else text


def _content_type_for(path: Path) -> str:
    if path.suffix == ".txt":
        return "text/plain"
    if path.suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if path.suffix == ".pdf":
        return "application/pdf"
    return "application/octet-stream"


def _percentile(values: list[float], percentile: int) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    index = round((len(ordered) - 1) * percentile / 100)
    return ordered[index]


def main() -> None:
    results = run_integration_test()
    summary = results["summary"]
    print(f"Uploaded documents: {summary['document_count']}")
    print(f"Evaluation questions: {summary['question_count']}")
    print(f"Retrieval accuracy: {summary['retrieval_accuracy']:.2%}")
    print(f"Answer accuracy: {summary['answer_accuracy']:.2%}")
    print(f"P95 response time: {summary['p95_response_ms']} ms")
    print(f"Report: {REPORT_PATH}")
    print(f"Raw results: {RESULTS_PATH}")


if __name__ == "__main__":
    main()
