from pathlib import Path
import sys

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document as DocxDocument

from scripts.day2.create_sample_pdf import build_pdf_bytes


SAMPLE_TEXT = (
    "Enterprise document parsing sample.\n"
    "Employees can search policies, product manuals, and FAQ files.\n"
    "Parsed documents must keep file metadata for later source citation.\n"
)


def create_pdf(path: Path) -> None:
    path.write_bytes(
        build_pdf_bytes(
            [
                "Enterprise document parsing sample.",
                "Employees can search policies, product manuals, and FAQ files.",
                "Parsed documents must keep file metadata for later source citation.",
            ]
        )
    )


def create_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_paragraph("Enterprise document parsing sample.")
    document.add_paragraph("Employees can search policies, product manuals, and FAQ files.")
    document.add_paragraph("Parsed documents must keep file metadata for later source citation.")
    document.save(path)


def create_txt(path: Path) -> None:
    path.write_text(SAMPLE_TEXT, encoding="utf-8")


def main() -> None:
    sample_dir = Path("data/day3_samples")
    sample_dir.mkdir(parents=True, exist_ok=True)

    create_pdf(sample_dir / "sample_policy.pdf")
    create_docx(sample_dir / "sample_manual.docx")
    create_txt(sample_dir / "sample_faq.txt")

    print(f"Created sample documents in {sample_dir}")


if __name__ == "__main__":
    main()
