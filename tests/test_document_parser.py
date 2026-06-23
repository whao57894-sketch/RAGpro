from pathlib import Path

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document

from scripts.day3.create_sample_documents import main as create_sample_documents
from src.document_parser import UnsupportedDocumentTypeError, parse_document, parse_documents


EXPECTED_PHRASES = [
    "Enterprise document parsing sample.",
    "Employees can search policies",
    "Parsed documents must keep file metadata",
]


def test_parse_pdf_docx_txt_documents():
    create_sample_documents()
    sample_dir = Path("data/day3_samples")
    paths = [
        sample_dir / "sample_policy.pdf",
        sample_dir / "sample_manual.docx",
        sample_dir / "sample_faq.txt",
    ]

    documents = parse_documents(paths)

    assert documents
    assert all(isinstance(document, Document) for document in documents)

    parsed_text = "\n".join(document.page_content for document in documents)
    for phrase in EXPECTED_PHRASES:
        assert phrase in parsed_text

    file_names = {document.metadata["file_name"] for document in documents}
    assert file_names == {"sample_policy.pdf", "sample_manual.docx", "sample_faq.txt"}

    for document in documents:
        assert Path(document.metadata["file_path"]).exists()
        assert document.metadata["file_extension"] in {".pdf", ".docx", ".txt"}
        assert isinstance(document.metadata["document_index"], int)


def test_parse_unsupported_file_type(tmp_path):
    file_path = tmp_path / "sample.xlsx"
    file_path.write_text("unsupported", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentTypeError):
        parse_document(file_path)


def test_parse_docx_keeps_heading_and_table_structure(tmp_path):
    file_path = tmp_path / "employee_profile.docx"
    doc = DocxDocument()
    doc.add_heading("项目基本信息", level=1)
    doc.add_paragraph("该文档用于登记项目成员和模块信息。")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "内容"
    table.rows[1].cells[0].text = "姓名"
    table.rows[1].cells[1].text = "张三"
    table.rows[2].cells[0].text = "项目名称"
    table.rows[2].cells[1].text = "企业知识库系统"
    doc.save(file_path)

    documents = parse_document(file_path)

    assert documents
    section_types = {document.metadata["section_type"] for document in documents}
    assert {"heading", "paragraph", "table_row"} <= section_types
    row_documents = [document for document in documents if document.metadata["section_type"] == "table_row"]
    assert any("姓名：张三" in document.page_content for document in row_documents)
    assert any("项目名称：企业知识库系统" in document.page_content for document in row_documents)
    assert all("项目基本信息" in document.page_content for document in row_documents)
    assert any("姓名" in document.metadata["field_keys"] for document in row_documents)
