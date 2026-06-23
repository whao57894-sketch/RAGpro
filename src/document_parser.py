import re
from pathlib import Path
from typing import Iterable, Iterator

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
FIELD_PATTERN = re.compile(r"([\u4e00-\u9fffA-Za-z0-9_()（）/\-\s]{1,24})[:：]")
DOCX_HEADING_PREFIXES = ("heading", "标题")


class UnsupportedDocumentTypeError(ValueError):
    """Raised when the parser receives an unsupported document type."""


def parse_document(file_path: str | Path) -> list[Document]:
    path = Path(file_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    if not path.is_file():
        raise ValueError(f"Document path is not a file: {path}")

    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {extension}")

    if extension == ".docx":
        return _parse_docx_document(path)

    loader = _build_loader(path, extension)
    documents = loader.load()
    return _normalize_documents(documents, path, extension)


def parse_documents(file_paths: Iterable[str | Path]) -> list[Document]:
    parsed: list[Document] = []
    for file_path in file_paths:
        parsed.extend(parse_document(file_path))
    return parsed


def _build_loader(path: Path, extension: str):
    if extension == ".pdf":
        return PyPDFLoader(str(path))
    if extension == ".txt":
        return TextLoader(str(path), encoding="utf-8")
    raise UnsupportedDocumentTypeError(f"Unsupported document type: {extension}")


def _parse_docx_document(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    heading_stack: list[str] = []
    normalized: list[Document] = []

    for document_index, block in enumerate(_iter_docx_blocks(doc)):
        if isinstance(block, Paragraph):
            text = _normalize_text(block.text)
            if not text:
                continue
            style_name = (block.style.name or "").strip()
            if _is_heading_style(style_name):
                heading_stack = _update_heading_stack(heading_stack, style_name, text)
                normalized.append(
                    _build_docx_document(
                        text=text,
                        path=path,
                        document_index=document_index,
                        section_type="heading",
                        heading_path=heading_stack,
                    )
                )
                continue

            section_type = "field_block" if _extract_field_keys(text) else "paragraph"
            normalized.append(
                _build_docx_document(
                    text=text,
                    path=path,
                    document_index=document_index,
                    section_type=section_type,
                    heading_path=heading_stack,
                )
            )
            continue

        if isinstance(block, Table):
            normalized.extend(
                _parse_docx_table(
                    block=block,
                    path=path,
                    base_index=document_index,
                    heading_path=heading_stack,
                )
            )

    return normalized


def _parse_docx_table(
    block: Table,
    path: Path,
    base_index: int,
    heading_path: list[str],
) -> list[Document]:
    rows = [
        [_normalize_text(cell.text) for cell in row.cells]
        for row in block.rows
    ]
    rows = [[cell for cell in row if cell] for row in rows]
    rows = [row for row in rows if row]
    if not rows:
        return []

    header_row = rows[0] if _looks_like_header_row(rows[0]) else None
    data_rows = rows[1:] if header_row else rows
    documents: list[Document] = []

    for row_index, row in enumerate(data_rows):
        row_text, field_keys = _table_row_to_text(row, header_row)
        if not row_text:
            continue
        documents.append(
            _build_docx_document(
                text=row_text,
                path=path,
                document_index=base_index * 1000 + row_index,
                section_type="table_row",
                heading_path=heading_path,
                field_keys=field_keys,
                table_row_index=row_index,
            )
        )
    return documents


def _table_row_to_text(row: list[str], header_row: list[str] | None) -> tuple[str, list[str]]:
    if header_row and len(header_row) == 2:
        normalized_headers = [_normalize_text(item) for item in header_row]
        if normalized_headers in (["字段", "内容"], ["字段", "值"], ["项目", "内容"]):
            field_key = _normalize_text(row[0]) if row else ""
            field_value = _normalize_text(row[1]) if len(row) > 1 else ""
            if field_key and field_value:
                return f"{field_key}：{field_value}", [field_key]

    if header_row and len(header_row) == len(row):
        pairs = []
        field_keys: list[str] = []
        for header, value in zip(header_row, row):
            if not value:
                continue
            key = _normalize_text(header)
            if key:
                field_keys.append(key)
                pairs.append(f"{key}：{value}")
            else:
                pairs.append(value)
        return "；".join(pairs), field_keys

    if len(row) == 2:
        field_key = _normalize_text(row[0])
        field_value = _normalize_text(row[1])
        if field_key and field_value:
            return f"{field_key}：{field_value}", [field_key]

    return "；".join(cell for cell in row if cell), _extract_field_keys("；".join(row))


def _build_docx_document(
    *,
    text: str,
    path: Path,
    document_index: int,
    section_type: str,
    heading_path: list[str],
    field_keys: list[str] | None = None,
    table_row_index: int | None = None,
) -> Document:
    raw_text = _normalize_text(text)
    heading_text = " > ".join(item for item in heading_path if item)
    retrieval_text = _build_retrieval_text(raw_text, heading_text, section_type, field_keys or [])
    metadata = {
        "file_name": path.name,
        "file_path": str(path),
        "file_extension": path.suffix.lower(),
        "document_index": document_index,
        "section_type": section_type,
        "heading_path": heading_text,
        "display_text": raw_text,
        "retrieval_text": retrieval_text,
        "field_keys": "|".join(field_keys or _extract_field_keys(raw_text)),
    }
    if table_row_index is not None:
        metadata["table_row_index"] = table_row_index
    return Document(page_content=retrieval_text, metadata=metadata)


def _normalize_documents(documents: list[Document], path: Path, extension: str) -> list[Document]:
    normalized: list[Document] = []
    for index, document in enumerate(documents):
        text = _normalize_text(document.page_content)
        metadata = {
            **document.metadata,
            "file_name": path.name,
            "file_path": str(path),
            "file_extension": extension,
            "document_index": index,
            "section_type": document.metadata.get("section_type", "paragraph"),
            "heading_path": document.metadata.get("heading_path", ""),
            "display_text": text,
            "retrieval_text": text,
            "field_keys": "|".join(_extract_field_keys(text)),
        }
        normalized.append(Document(page_content=text, metadata=metadata))
    return normalized


def _iter_docx_blocks(doc: DocxDocumentType) -> Iterator[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _is_heading_style(style_name: str) -> bool:
    normalized = style_name.strip().lower()
    return normalized.startswith(DOCX_HEADING_PREFIXES)


def _update_heading_stack(current: list[str], style_name: str, title: str) -> list[str]:
    match = re.search(r"(\d+)$", style_name)
    if match:
        level = max(int(match.group(1)), 1)
    else:
        level = 1
    next_stack = current[: level - 1]
    next_stack.append(title)
    return next_stack


def _looks_like_header_row(row: list[str]) -> bool:
    if len(row) < 2:
        return False
    normalized = [_normalize_text(cell) for cell in row if cell]
    if len(normalized) < 2:
        return False
    return len(set(normalized)) == len(normalized) and all(len(cell) <= 20 for cell in normalized)


def _build_retrieval_text(text: str, heading_text: str, section_type: str, field_keys: list[str]) -> str:
    lines = []
    if heading_text and section_type != "heading":
        lines.append(f"章节：{heading_text}")
    if field_keys:
        lines.append(f"字段：{'、'.join(dict.fromkeys(field_keys))}")
    lines.append(text)
    return "\n".join(line for line in lines if line)


def _extract_field_keys(text: str) -> list[str]:
    keys = []
    for match in FIELD_PATTERN.findall(text):
        key = _normalize_text(match).strip("：:")
        if 0 < len(key) <= 24:
            keys.append(key)
    return list(dict.fromkeys(keys))


def _normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\t", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \u3000]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
